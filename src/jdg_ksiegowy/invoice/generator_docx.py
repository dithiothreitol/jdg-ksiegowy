"""Generator faktur w formacie DOCX (python-docx)."""

from __future__ import annotations

from datetime import date
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from jdg_ksiegowy.config import settings
from jdg_ksiegowy.invoice.models import Invoice


def _format_pln(amount: Decimal) -> str:
    """Formatuj kwote jako PLN."""
    return f"{amount:,.2f} PLN".replace(",", " ")


def _format_date(d: date) -> str:
    return d.strftime("%d.%m.%Y")


def _set_cell_text(cell, text: str, bold: bool = False, align: str = "left", size: int = 9):
    """Ustaw tekst komorki tabeli."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def generate_invoice_docx(invoice: Invoice, output_path: Path) -> Path:
    """Wygeneruj fakture DOCX na podstawie modelu Invoice."""
    seller = settings.seller
    doc = Document()

    # --- Styl dokumentu ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # --- Naglowek: FAKTURA VAT ---
    heading = doc.add_heading(f"FAKTURA VAT nr {invoice.number}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    # --- Daty ---
    dates_para = doc.add_paragraph()
    dates_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dates_para.add_run(f"Data wystawienia: {_format_date(invoice.issue_date)}\n").font.size = Pt(9)
    dates_para.add_run(f"Data sprzedazy: {_format_date(invoice.sale_date)}\n").font.size = Pt(9)
    dates_para.add_run(
        f"Termin platnosci: {_format_date(invoice.payment_due)}"
    ).font.size = Pt(9)

    if invoice.period_from and invoice.period_to:
        dates_para.add_run(
            f"\nOkres: {_format_date(invoice.period_from)} - {_format_date(invoice.period_to)}"
        ).font.size = Pt(9)

    # --- Sprzedawca / Nabywca (tabela 2 kolumny) ---
    parties_table = doc.add_table(rows=1, cols=2)
    parties_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Sprzedawca
    seller_cell = parties_table.cell(0, 0)
    seller_cell.text = ""
    p = seller_cell.paragraphs[0]
    p.add_run("SPRZEDAWCA\n").bold = True
    p.add_run(f"{seller.name}\n")
    p.add_run(f"{seller.address}\n")
    p.add_run(f"NIP: {seller.nip}\n")
    p.add_run(f"Bank: {seller.bank_name}\n")
    p.add_run(f"Nr konta: {seller.bank_account}")

    # Nabywca
    buyer_cell = parties_table.cell(0, 1)
    buyer_cell.text = ""
    p = buyer_cell.paragraphs[0]
    p.add_run("NABYWCA\n").bold = True
    p.add_run(f"{invoice.buyer.name}\n")
    p.add_run(f"{invoice.buyer.address}\n")
    p.add_run(f"NIP: {invoice.buyer.nip}")

    doc.add_paragraph()  # spacer

    # --- Tabela pozycji ---
    headers = ["Lp", "Opis", "Ilosc", "J.m.", "Cena netto", "Wart. netto", "VAT %", "Kwota VAT", "Wart. brutto"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Naglowki
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, align="center", size=8)

    # Pozycje
    for idx, item in enumerate(invoice.items, 1):
        row = table.add_row()
        cells = row.cells
        _set_cell_text(cells[0], str(idx), align="center")
        _set_cell_text(cells[1], item.description)
        _set_cell_text(cells[2], str(item.quantity), align="center")
        _set_cell_text(cells[3], item.unit, align="center")
        _set_cell_text(cells[4], f"{item.unit_price_net:.2f}", align="right")
        _set_cell_text(cells[5], f"{item.net_value:.2f}", align="right")
        _set_cell_text(cells[6], f"{item.vat_rate}%", align="center")
        _set_cell_text(cells[7], f"{item.vat_amount:.2f}", align="right")
        _set_cell_text(cells[8], f"{item.gross_value:.2f}", align="right")

    # Wiersz podsumowania
    summary_row = table.add_row()
    _set_cell_text(summary_row.cells[0], "")
    _set_cell_text(summary_row.cells[1], "RAZEM", bold=True, align="right")
    _set_cell_text(summary_row.cells[2], "")
    _set_cell_text(summary_row.cells[3], "")
    _set_cell_text(summary_row.cells[4], "")
    _set_cell_text(summary_row.cells[5], f"{invoice.total_net:.2f}", bold=True, align="right")
    _set_cell_text(summary_row.cells[6], "")
    _set_cell_text(summary_row.cells[7], f"{invoice.total_vat:.2f}", bold=True, align="right")
    _set_cell_text(summary_row.cells[8], f"{invoice.total_gross:.2f}", bold=True, align="right")

    # Ustawienia szerokosc kolumn
    widths = [Cm(1), Cm(6), Cm(1.2), Cm(1.2), Cm(2.2), Cm(2.2), Cm(1.2), Cm(2.2), Cm(2.5)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    doc.add_paragraph()  # spacer

    # --- Podsumowanie kwot ---
    summary = doc.add_paragraph()
    summary.add_run(f"Do zaplaty: {_format_pln(invoice.total_gross)}\n").bold = True
    summary.add_run(f"Forma platnosci: przelew bankowy\n")
    summary.add_run(f"Nr konta: {seller.bank_account}\n")
    summary.add_run(f"Termin platnosci: {_format_date(invoice.payment_due)}")

    if invoice.notes:
        doc.add_paragraph()
        doc.add_paragraph(f"Uwagi: {invoice.notes}")

    # --- Stopka ---
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Dokument wygenerowany elektronicznie — nie wymaga podpisu")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Zapisz
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
